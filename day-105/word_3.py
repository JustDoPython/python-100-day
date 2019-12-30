# word_3.py

# 引入库
from docx import Document

# 打开文档1
doc1 = Document('word1.docx')

# 读取每段内容
pl = [ paragraph.text for paragraph in doc1.paragraphs]

print('###### 输出word1文章内容')
# 输出读取到的内容
for i in pl:
    print(i)


# 打开文档2
doc2 = Document('word2.docx')

print('\n###### 输出word2文章内容')

pl2 = [ paragraph.text for paragraph in doc2.paragraphs]

# 输出读取到的内容
for j in pl2:
    print(j)

# 读取表格材料，并输出结果
tables = [table for table in doc2.tables]
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            print (cell.text,end='  ')
        print()
    print('\n')
    