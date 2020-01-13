# word_1.py

# 导入库
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor

# 新建文档
doc1 = Document()

# 新增文档标题
doc1.add_heading('如何使用 Python 创建 Word',0)

# 创建段落描述
doc1.add_paragraph('我们平时使用 Word 用来做文章的处理，可能没想过它可以用 Python 生成，下面我们就介绍具体如何操作……')

# 创建一级标题
doc1.add_heading('安装 python-docx 库',1)

# 创建段落描述
doc1.add_paragraph('现在开始我们来介绍如何安装 python-docx 库，具体需要以下两步操作：')

# 创建二级标题
doc1.add_heading('第一步：安装 Python',2)

# 创建段落，添加文档内容
paragraph = doc1.add_paragraph('这是第一步的安装描述！')

# 段落中增加文字，并设置字体字号
run = paragraph.add_run('(注意：这里设置了字号为20)')
run.font.size = Pt(20)

# 设置英文字体
run = doc1.add_paragraph('这里设置英文字体：').add_run('This Font is Times New Roman ')
run.font.name = 'Times New Roman'

# 设置中文字体
run = doc1.add_paragraph('这里设置中文字体：').add_run('当前字体为黑体')
run.font.name='黑体'
ele = run._element
ele.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 设置斜体
run = doc1.add_paragraph('这段设置：').add_run('文字的是斜体 ')
run.italic = True

# 设置粗体
run = doc1.add_paragraph('这段再设置：').add_run('这里设置粗体').bold = True

run = doc1.add_paragraph('这段为下划线：').add_run('这里设置带下划线').underline = True

run = doc1.add_paragraph('这段字体为红色：').add_run('这里设置字体为红色')
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# 增加引用
doc1.add_paragraph('这里是我们引用的一段话：人生苦短，我用Python。', style='Intense Quote')

# 保存文件
doc1.save('word1.docx')